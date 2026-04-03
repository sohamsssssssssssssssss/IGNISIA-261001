"""
Versioned GSTIN-based MSME credit scoring API.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, Literal, Optional

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from ..core.entity_graph import get_entity_graph_service
from ..schemas.score import SimulationResponse
from ..core.gstin import is_valid_gstin, normalize_gstin
from ..core.security import require_role
from ..core.settings import get_settings
from ..core.storage import get_storage
from ..core.xgboost_model import (
    REAL_OUTCOME_MIN_RECORDS,
    get_risk_band,
    get_scorer,
    recommend_loan,
    resolve_industry_profile,
)
from ..services.feature_engineering import build_feature_vector
from ..services.graph_serializer import serialize_graph
from ..services.simulator import run_simulation
from ..core.scheduler import refresh_gstin_now, refresh_pipeline_stream, trigger_immediate_ingestion
from ..services.upi_fraud_detection import UPIFraudDetector
from ..utils.audit import AuditTrail


router = APIRouter(prefix="/api", tags=["MSME Credit Scoring"])
logger = logging.getLogger("intellicredit.scoring")

class ModelRetrainResponse(BaseModel):
    status: str
    model: Dict[str, Any]
    retrained: bool
    outcomes_recorded: int
    total_real_outcomes: int
    governance: Dict[str, Any]


class LoanOutcomePayload(BaseModel):
    gstin: str
    outcome: Literal["repaid", "defaulted"]
    loan_amount: float
    tenure_months: int
    company_name: Optional[str] = None


class ModelRetrainRequest(BaseModel):
    outcomes: list[LoanOutcomePayload] = []


FRESH_MAX_MINUTES = 30
EXPIRED_AFTER_MINUTES = 240


def _classify_outcome(credit_score: int, data_sparse: bool) -> str:
    """Derive a human-readable scenario label from the model's actual output."""
    if data_sparse:
        return "sparse"
    if credit_score >= 650:
        return "approve"
    return "reject"

def _build_confidence_summary(features: Dict[str, float]) -> Dict[str, Any]:
    return {
        "overall_data_confidence": round(features.get("overall_data_confidence", 0.0), 2),
        "gst_confidence": round(features.get("gst_data_confidence", 0.0), 2),
        "upi_confidence": round(features.get("upi_data_confidence", 0.0), 2),
        "eway_confidence": round(features.get("eway_data_confidence", 0.0), 2),
    }


def _build_data_sources(pipeline_data: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "source_mode": "mocked",
        "judge_note": "All GST, UPI, and e-way bill feeds in this environment are simulated for demo use.",
        "pipelines": {
            "gst_velocity": {
                "source_status": "mock",
                "description": "Synthetic GST filing timeliness and e-invoice velocity stream",
                "freshness": pipeline_data["gst_velocity"]["data_freshness"],
            },
            "upi_cadence": {
                "source_status": "mock",
                "description": "Synthetic UPI cadence and transaction topology stream",
                "freshness": pipeline_data["upi_cadence"]["data_freshness"],
            },
            "eway_bill": {
                "source_status": "mock",
                "description": "Synthetic e-way bill volume and anomaly stream",
                "freshness": pipeline_data["eway_bill"]["data_freshness"],
            },
        },
    }

def _build_data_freshness(pipeline_data: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "gst": {
            "stream": "GST Filing",
            "fetched_at": pipeline_data["gst_velocity"]["data_freshness"],
        },
        "upi": {
            "stream": "UPI Transactions",
            "fetched_at": pipeline_data["upi_cadence"]["data_freshness"],
        },
        "eway": {
            "stream": "E-Way Bill",
            "fetched_at": pipeline_data["eway_bill"]["data_freshness"],
        },
    }


def _derive_missing_streams(pipeline_data: Dict[str, Any]) -> list[str]:
    stream_map = {
        "gst": pipeline_data["gst_velocity"],
        "upi": pipeline_data["upi_cadence"],
        "eway": pipeline_data["eway_bill"],
    }
    return [name for name, payload in stream_map.items() if payload.get("sparse_data")]

def _build_score_simulation(
    gstin: str,
    company_name: Optional[str] = None,
    *,
    request_id: Optional[str] = None,
) -> Dict[str, Any]:
    base_payload = _score_assessment_payload(
        gstin,
        company_name,
        request_id=request_id,
        persist=False,
    )
    feature_vector = dict(base_payload["feature_vector"])
    shap_values = {
        item["feature_key"]: float(item["shap_value"])
        for item in base_payload["all_shap_values"]
        if item.get("feature_key")
    }

    def _score_features(vector: Dict[str, Any]) -> int:
        return int(get_scorer().score(vector)["credit_score"])

    return run_simulation(
        gstin=base_payload["gstin"],
        feature_vector=feature_vector,
        shap_values=shap_values,
        scorer_fn=_score_features,
        approval_threshold=550,
    )


def _parse_iso_timestamp(value: str) -> datetime:
    return datetime.fromisoformat(value.replace("Z", "+00:00"))


def _compute_freshness(pipeline_data: Dict[str, Any]) -> Dict[str, Any]:
    timestamps = pipeline_data.get("pipeline_ingested_at", {})
    if not timestamps:
        oldest_timestamp = pipeline_data.get("pipeline_timestamp")
        if oldest_timestamp is None:
            return {
                "oldest_timestamp": None,
                "staleness_minutes": None,
                "status": "unknown",
                "manual_review_required": False,
            }
        oldest_dt = _parse_iso_timestamp(oldest_timestamp)
    else:
        oldest_timestamp = min(timestamps.values())
        oldest_dt = _parse_iso_timestamp(oldest_timestamp)

    now = datetime.now(timezone.utc)
    staleness_minutes = max(0, int((now - oldest_dt).total_seconds() // 60))
    if staleness_minutes > EXPIRED_AFTER_MINUTES:
        status = "expired"
        manual_review_required = False
    elif staleness_minutes >= FRESH_MAX_MINUTES:
        status = "stale"
        manual_review_required = True
    else:
        status = "fresh"
        manual_review_required = False

    return {
        "oldest_timestamp": oldest_timestamp,
        "staleness_minutes": staleness_minutes,
        "status": status,
        "manual_review_required": manual_review_required,
    }


def _build_outcome_feature_snapshot(gstin: str) -> Dict[str, Any]:
    storage = get_storage()
    pipeline_data = storage.get_pipeline_data(gstin)
    if pipeline_data is None:
        trigger_immediate_ingestion(gstin)
        raise HTTPException(
            status_code=202,
            detail="Feature snapshot unavailable for this GSTIN. Ingestion has been triggered; retry outcome submission shortly.",
            headers={"Retry-After": str(min(get_settings().pipeline_interval_seconds, 30))},
        )
    return build_feature_vector(pipeline_data)


def _score_assessment_payload(
    gstin: str,
    company_name: Optional[str] = None,
    industry_code: Optional[str] = None,
    *,
    request_id: Optional[str] = None,
    persist: bool = True,
) -> Dict[str, Any]:
    settings = get_settings()
    started_at = datetime.now(timezone.utc)
    audit_trail = AuditTrail()
    audit_log = audit_trail.log
    normalized_gstin = normalize_gstin(gstin)
    if not is_valid_gstin(normalized_gstin):
        raise HTTPException(
            status_code=400,
            detail="Invalid GSTIN format. Expected 15 characters: state code + PAN + entity code + Z + checksum.",
        )
    audit_log(f"GSTIN received: {normalized_gstin}")

    # Read from the persistent pipeline data store.
    # If data hasn't been ingested yet, register the GSTIN and let background
    # workers populate the store out of band.
    storage = get_storage()
    pipeline_data = storage.get_pipeline_data(normalized_gstin)
    if pipeline_data is None:
        audit_log("No pipeline snapshot found — background ingestion scheduled")
        trigger_immediate_ingestion(normalized_gstin, company_name)
        retry_after = min(get_settings().pipeline_interval_seconds, 30)
        raise HTTPException(
            status_code=202,
            detail="Data ingestion in progress for this GSTIN. Retry after background workers ingest the first pipeline snapshots.",
            headers={"Retry-After": str(retry_after)},
        )
    audit_log("Mock pipelines initialised — GST, UPI, e-Way Bill")
    freshness = _compute_freshness(pipeline_data)
    if freshness["status"] == "expired":
        trigger_immediate_ingestion(normalized_gstin, company_name)
        retry_after = min(get_settings().pipeline_interval_seconds, 300)
        raise HTTPException(
            status_code=202,
            detail=(
                f"Underlying pipeline data is expired ({freshness['staleness_minutes']} minutes old). "
                "Re-ingestion has been triggered; retry after fresh pipeline snapshots are available."
            ),
            headers={"Retry-After": str(retry_after)},
        )
    audit_log("Mock pipelines initialised — GST, UPI, e-Way Bill")
    features = build_feature_vector(pipeline_data)
    audit_log(f"Feature engineering complete — {len(features)} features extracted")

    scorer = get_scorer()
    score_result = scorer.score(features)
    audit_log(f"Gradient boosting scored — raw probability: {score_result['probability']:.3f}")
    audit_log(f"Score mapped to 300-900 scale: {score_result['credit_score']}")
    if score_result["top_reasons"]:
        top_negative = next(
            (reason for reason in score_result["top_reasons"] if reason.get("shap_value", 0) < 0),
            score_result["top_reasons"][0],
        )
        audit_log(
            f"SHAP computed — top driver: {top_negative['feature']} ({top_negative['shap_value']:+.4f})"
        )

    fraud_detector = UPIFraudDetector()
    upi_transactions = pipeline_data.get("upi_cadence", {}).get("transactions", [])
    audit_log("Fraud detector running — UPI transaction graph building")
    fraud_result = fraud_detector.detect_circular_transactions(
        upi_transactions,
        normalized_gstin,
        entity_graph_service=get_entity_graph_service(),
    )
    if fraud_result["cycle_count"] > 0:
        audit_log(
            f"NetworkX cycle detected — {fraud_result['cycle_count']} cycles, ₹{fraud_result['total_volume']:,} traced"
        )
    else:
        audit_log("No circular UPI topology detected")

    if fraud_result["circular_risk"] == "HIGH":
        score_result["credit_score"] = max(300, score_result["credit_score"] - 150)
        score_result["fraud_penalty_applied"] = True
        score_result["top_reasons"].insert(
            0,
            {
                "feature": "UPI Circular Transaction Detection",
                "shap_value": -0.5,
                "feature_value": fraud_result["risk_score"],
                "direction": "negative",
                "reason": f"Circular fund rotation detected across {fraud_result['cycle_count']} cycles — score penalized by 150 points",
            },
        )
        audit_log(f"Circular topology penalty applied: -150 points -> final score {score_result['credit_score']}")
    else:
        score_result["fraud_penalty_applied"] = False

    score_result["risk_band"] = get_risk_band(score_result["credit_score"])
    audit_log(f"Risk band assigned: {score_result['risk_band']['band']}")
    score_result["recommendation"] = recommend_loan(
        score_result["credit_score"],
        abs(features.get("upi_net_cash_flow", 500000)),
        fraud_risk=fraud_result["circular_risk"],
        industry_code=industry_code,
        data_confidence=features.get("overall_data_confidence", 1.0),
        months_active=features.get("history_months_active", 12.0),
    )
    score_result["freshness_status"] = freshness["status"]
    score_result["manual_review_required"] = freshness["manual_review_required"]
    if freshness["manual_review_required"]:
        score_result["recommendation"]["manual_review_required"] = True
        score_result["recommendation"]["manual_review_reason"] = (
            f"Underlying pipeline data is {freshness['staleness_minutes']} minutes old; "
            "manual review is recommended before an automated credit decision."
        )
    else:
        score_result["recommendation"]["manual_review_required"] = False

    # Derive scenario label post-hoc from the model's actual output
    data_sparse = any(
        [
            pipeline_data["gst_velocity"]["sparse_data"],
            pipeline_data["upi_cadence"]["sparse_data"],
            pipeline_data["eway_bill"]["sparse_data"],
        ]
    )
    if data_sparse:
        audit_log(
            f"Sparse data detected — confidence weight: {features.get('overall_data_confidence', 0.0):.2f}"
        )
    scenario = _classify_outcome(score_result["credit_score"], data_sparse)

    history_entry_timestamp = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    effective_freshness_timestamp = freshness["oldest_timestamp"] or score_result["score_freshness"]
    display_company_name = company_name or f"MSME Entity ({normalized_gstin[:5]}...)"
    if persist:
        storage.record_assessment(
            gstin=normalized_gstin,
            company_name=display_company_name,
            credit_score=score_result["credit_score"],
            risk_band=score_result["risk_band"]["band"],
            fraud_risk=fraud_result["circular_risk"],
            model_version=score_result["model_version"],
            industry_code=("".join(ch for ch in (industry_code or "") if ch.isdigit())[:2] or None),
            months_active=float(features.get("history_months_active", 0.0)),
            scenario=scenario,
            data_sparse=data_sparse,
            freshness_timestamp=effective_freshness_timestamp,
            created_at=history_entry_timestamp,
            top_reasons=score_result["top_reasons"][:5],
            recommendation=score_result["recommendation"],
        )

        if fraud_result["circular_risk"] in {"HIGH", "MEDIUM"}:
            storage.record_fraud_alert(
                gstin=normalized_gstin,
                company_name=display_company_name,
                circular_risk=fraud_result["circular_risk"],
                risk_score=fraud_result["risk_score"],
                cycle_count=fraud_result["cycle_count"],
                linked_msme_count=fraud_result.get("linked_msme_count", 0),
                total_volume=fraud_result["total_volume"],
                created_at=history_entry_timestamp,
                payload=fraud_result,
            )
    audit_log("Scorecard payload compiled — recommendation and history ready")
    elapsed_seconds = (datetime.now(timezone.utc) - started_at).total_seconds()
    audit_log(f"Complete — total processing time: {elapsed_seconds:.2f}s")

    history_page_size = min(settings.default_history_page_size, settings.max_history_page_size)
    score_history = storage.get_score_history(
        normalized_gstin,
        limit=history_page_size,
    )
    percentile = storage.get_segment_score_percentile(
        credit_score=score_result["credit_score"],
        months_active=float(features.get("history_months_active", 0.0)),
        industry_code=industry_code,
    )
    confidence_summary = _build_confidence_summary(features)
    data_sources = _build_data_sources(pipeline_data)
    data_freshness = _build_data_freshness(pipeline_data)
    missing_streams = _derive_missing_streams(pipeline_data)

    payload = {
        "gstin": normalized_gstin,
        "company_name": display_company_name,
        "industry_code": industry_code,
        "industry_profile": resolve_industry_profile(industry_code),
        "base_score": score_result.get("base_score"),
        "credit_score": score_result["credit_score"],
        "risk_band": score_result["risk_band"],
        "top_reasons": score_result["top_reasons"][:5],
        "all_shap_values": score_result["all_shap_values"],
        "recommendation": score_result["recommendation"],
        "fraud_detection": fraud_result,
        "fraud_ring_members": fraud_result.get("fraud_ring_members", []),
        "cross_entity_fraud_detected": fraud_result.get("cross_entity_fraud_detected", False),
        "fraud_penalty_applied": score_result.get("fraud_penalty_applied", False),
        "pipeline_signals": {
            "gst_velocity": {
                "filing_rate": pipeline_data["gst_velocity"]["velocity_metrics"]["filings_per_month"],
                "avg_delay": pipeline_data["gst_velocity"]["velocity_metrics"]["avg_delay_days"],
                "on_time_pct": pipeline_data["gst_velocity"]["velocity_metrics"]["on_time_pct"],
                "e_invoice_trend": pipeline_data["gst_velocity"]["velocity_metrics"]["e_invoice_trend"],
                "months_active": pipeline_data["gst_velocity"]["months_active"],
                "sparse_data": pipeline_data["gst_velocity"]["sparse_data"],
                "confidence_weight": confidence_summary["gst_confidence"],
                "data_freshness": pipeline_data["gst_velocity"]["data_freshness"],
            },
            "upi_cadence": {
                "avg_daily_txns": pipeline_data["upi_cadence"]["cadence_metrics"]["avg_daily_txns"],
                "regularity_score": pipeline_data["upi_cadence"]["cadence_metrics"]["regularity_score"],
                "inflow_outflow_ratio": pipeline_data["upi_cadence"]["flow_pattern"]["inflow_outflow_ratio"],
                "round_amount_pct": pipeline_data["upi_cadence"]["flow_pattern"]["round_amount_pct"],
                "months_active": pipeline_data["upi_cadence"]["months_active"],
                "sparse_data": pipeline_data["upi_cadence"]["sparse_data"],
                "confidence_weight": confidence_summary["upi_confidence"],
                "data_freshness": pipeline_data["upi_cadence"]["data_freshness"],
            },
            "eway_bill": {
                "avg_monthly_bills": pipeline_data["eway_bill"]["trend_metrics"]["avg_bills_per_month"],
                "volume_momentum": pipeline_data["eway_bill"]["trend_metrics"]["volume_momentum_pct"],
                "interstate_ratio": pipeline_data["eway_bill"]["trend_metrics"]["interstate_ratio"],
                "anomaly_count": len(pipeline_data["eway_bill"]["anomaly_flags"]),
                "months_active": pipeline_data["eway_bill"]["months_active"],
                "sparse_data": pipeline_data["eway_bill"]["sparse_data"],
                "confidence_weight": confidence_summary["eway_confidence"],
                "data_freshness": pipeline_data["eway_bill"]["data_freshness"],
            },
        },
        "feature_vector": features,
        "confidence_summary": confidence_summary,
        "data_sources": data_sources,
        "data_freshness": data_freshness,
        "score_history": score_history,
        "score_freshness": effective_freshness_timestamp,
        "data_ingested_at": effective_freshness_timestamp,
        "data_staleness_minutes": freshness["staleness_minutes"],
        "freshness_status": freshness["status"],
        "manual_review_required": freshness["manual_review_required"],
        "model_inference_at": score_result["score_freshness"],
        "model_version": score_result["model_version"],
        "model_metrics": score_result.get("model_metrics", {}),
        "model_backend": scorer.health_summary().get("backend"),
        "percentile": percentile,
        "calibration": score_result.get("calibration", {}),
        "calibration_method": score_result.get("calibration_method"),
        "score_mapping": score_result.get("score_mapping"),
        "probability": score_result.get("probability"),
        "raw_probability": score_result.get("raw_probability"),
        "default_probability": score_result.get("default_probability"),
        "scenario": scenario,
        "data_sparse": data_sparse,
        "sparse_data": data_sparse,
        "missing_streams": missing_streams,
        "audit_trail": audit_trail.to_list(),
    }

    logger.info(
        json.dumps(
            {
                "event": "score_decision",
                "request_id": request_id,
                "gstin": normalized_gstin,
                "credit_score": payload["credit_score"],
                "risk_band": payload["risk_band"]["band"],
                "fraud_risk": fraud_result["circular_risk"],
                "model_version": payload["model_version"],
                "feature_vector": payload["feature_vector"],
                "confidence_summary": confidence_summary,
            }
        )
    )
    if persist:
        embedding_service = get_embedding_service()
        try:
            embedding_service.embed_score_payload(normalized_gstin, payload)
            if fraud_result.get("cycle_count", 0) > 0:
                fraud_doc_id = f"fraud:{normalized_gstin}:{payload['model_inference_at']}"
                fraud_text = "\n".join(
                    [
                        f"GSTIN: {normalized_gstin}",
                        f"Company: {display_company_name}",
                        f"Circular Risk: {fraud_result.get('circular_risk')}",
                        f"Cycle Count: {fraud_result.get('cycle_count')}",
                        f"Linked MSMEs: {fraud_result.get('linked_msme_count')}",
                        f"Fraud Ring Members: {', '.join(fraud_result.get('fraud_ring_members', [])) or 'none'}",
                        f"Cross Entity Fraud Detected: {fraud_result.get('cross_entity_fraud_detected')}",
                        f"Circular Flow Ratio: {fraud_result.get('circular_flow_ratio')}",
                        f"Total Volume: {fraud_result.get('total_volume')}",
                    ]
                )
                embedding_service.embed_document(
                    fraud_doc_id,
                    fraud_text,
                    {
                        "collection": "fraud_patterns",
                        "doc_type": "fraud_pattern",
                        "gstin": normalized_gstin,
                        "company_name": display_company_name,
                        "circular_risk": fraud_result.get("circular_risk"),
                        "cycle_count": fraud_result.get("cycle_count"),
                        "linked_msme_count": fraud_result.get("linked_msme_count"),
                        "score_freshness": payload["score_freshness"],
                    },
                )
        except Exception:
            logger.exception("Embedding pipeline failed for GSTIN %s", normalized_gstin)
        try:
            get_apriori_service().trigger_refresh_async_if_needed()
        except Exception:
            logger.exception("Apriori refresh trigger failed for GSTIN %s", normalized_gstin)

    if include_narrative:
        try:
            narrative_result = get_llm_service().generate_narrative(normalized_gstin, payload)
            payload["narrative"] = narrative_result.get("narrative_sections") or narrative_result["narrative"]
            payload["narrative_text"] = narrative_result["narrative_text"]
            payload["narrative_sources"] = narrative_result["sources"]
            payload["sources"] = narrative_result["sources"]
            payload["narrative_model_used"] = narrative_result["model_used"]
            if persist:
                storage.update_assessment_narrative(
                    gstin=normalized_gstin,
                    created_at=history_entry_timestamp,
                    narrative=payload["narrative_text"],
                )
        except Exception:
            logger.exception("Narrative generation failed for GSTIN %s", normalized_gstin)
            payload["narrative"] = None
            payload["narrative_text"] = None
            payload["narrative_sources"] = []
            payload["narrative_model_used"] = None
    return payload


def _render_docx(payload: Dict[str, Any]) -> BytesIO:
    doc = Document()
    doc.add_heading("MSME Credit Scorecard", 0)
    doc.add_paragraph(f"Company: {payload['company_name']}")
    doc.add_paragraph(f"GSTIN: {payload['gstin']}")
    doc.add_paragraph(f"Credit Score: {payload['credit_score']} / 900")
    doc.add_paragraph(f"Risk Band: {payload['risk_band']['band']}")
    doc.add_paragraph(f"Freshness: {payload['score_freshness']}")
    doc.add_paragraph(f"Model: {payload['model_version']}")
    doc.add_paragraph(f"Industry Code: {payload.get('industry_code') or 'Not provided'}")

    doc.add_heading("Confidence Summary", level=1)
    for key, value in payload["confidence_summary"].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Top Reasons", level=1)
    for idx, reason in enumerate(payload["top_reasons"], start=1):
        doc.add_paragraph(
            f"{idx}. {reason['feature']} — {reason['reason']} (SHAP {reason['shap_value']:+.4f})"
        )

    doc.add_heading("Loan Recommendation", level=1)
    recommendation = payload["recommendation"]
    for key in [
        "eligible",
        "recommended_amount",
        "recommended_tenure_months",
        "tenure_options_months",
        "indicative_rate_pct",
        "base_rate",
        "risk_premium",
        "repayment_structure",
        "recommendation_basis",
        "reason",
    ]:
        if key in recommendation:
            doc.add_paragraph(f"{key}: {recommendation[key]}")

    doc.add_heading("Feature Vector", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Feature"
    table.rows[0].cells[1].text = "Value"
    for key, value in sorted(payload["feature_vector"].items()):
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


@router.post(
    "/score/{gstin}",
    summary="Score an MSME by GSTIN",
    description="Returns a 300-900 credit score, risk band, SHAP reasons, confidence weights, recommendation, fraud assessment, and recent score history.",
)
@router.post(
    "/v1/score/{gstin}",
    summary="Score an MSME by GSTIN (v1)",
    description="Versioned scoring endpoint returning explainable MSME credit assessment from GST, UPI, and e-way bill mock pipelines.",
)
async def score_msme(
    gstin: str,
    request: Request,
    company_name: Optional[str] = Query(None, description="Optional company name for export and display"),
    industry_code: Optional[str] = Query(None, description="Optional NIC industry code for industry-aware recommendation logic"),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    return _score_assessment_payload(
        gstin,
        company_name,
        industry_code,
        request_id=getattr(request.state, "request_id", None),
        persist=True,
    )


@router.post(
    "/narrative",
    summary="Generate an underwriting narrative for the latest GSTIN score",
)
@router.post(
    "/v1/narrative",
    summary="Generate an underwriting narrative for the latest GSTIN score (v1)",
)
async def generate_narrative_endpoint(
    body: NarrativeRequest,
    request: Request,
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    payload = _load_latest_payload_for_llm(
        body.gstin,
        body.company_name,
        body.industry_code,
        request_id=getattr(request.state, "request_id", None),
    )
    result = get_llm_service().generate_narrative(payload["gstin"], payload)
    return {
        "gstin": payload["gstin"],
        "narrative": result.get("narrative_sections") or result["narrative"],
        "narrative_text": result.get("narrative_text"),
        "sources": result["sources"],
        "model_used": result["model_used"],
    }


@router.post(
    "/chat",
    summary="Ask a chat question about a scored GSTIN",
)
@router.post(
    "/v1/chat",
    summary="Ask a chat question about a scored GSTIN (v1)",
)
async def chat_endpoint(
    body: ChatRequest,
    request: Request,
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    payload = _load_latest_payload_for_llm(
        body.gstin,
        body.company_name,
        body.industry_code,
        request_id=getattr(request.state, "request_id", None),
    )
    result = get_llm_service().chat(
        payload["gstin"],
        payload,
        body.message,
        body.session_id,
    )
    return {
        "gstin": payload["gstin"],
        "reply": result["reply"],
        "sources": result["sources"],
        "session_id": result["session_id"],
        "sessionId": result["sessionId"],
        "model_used": result["model_used"],
    }


@router.get(
    "/insights/rules",
    summary="Return mined Apriori lending rules",
)
@router.get(
    "/v1/insights/rules",
    summary="Return mined Apriori lending rules (v1)",
)
async def get_mined_rules(
    force_refresh: bool = Query(False, description="Force recomputation instead of using the 1-hour cache"),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    service = get_apriori_service()
    rules = service.get_rules(force_refresh=force_refresh or service.rules_are_stale())
    return {
        "rules": rules,
        "cache": service.get_cache_metadata(),
    }
@router.get(
    "/score/{gstin}/history",
    summary="Fetch score history for a GSTIN",
)
@router.get(
    "/v1/score/{gstin}/history",
    summary="Fetch score history for a GSTIN (v1)",
)
async def get_score_history(
    gstin: str,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    normalized_gstin = normalize_gstin(gstin)
    if not is_valid_gstin(normalized_gstin):
        raise HTTPException(status_code=400, detail="Invalid GSTIN format")

    settings = get_settings()
    safe_page_size = min(page_size, settings.max_history_page_size)
    offset = (page - 1) * safe_page_size
    storage = get_storage()
    history = storage.get_score_history(
        normalized_gstin,
        limit=safe_page_size,
        offset=offset,
    )
    total = storage.get_assessment_count(normalized_gstin)
    return {
        "gstin": normalized_gstin,
        "history": history,
        "has_history": len(history) > 0,
        "page": page,
        "page_size": safe_page_size,
        "total_assessments": total,
        "has_more": offset + len(history) < total,
    }


@router.get(
    "/score/{gstin}/simulate",
    summary="Project score improvement over 6 months",
    response_model=SimulationResponse,
)
@router.get(
    "/v1/score/{gstin}/simulate",
    summary="Project score improvement over 6 months (v1)",
    response_model=SimulationResponse,
)
async def simulate_score_improvement(
    gstin: str,
    request: Request,
    company_name: Optional[str] = Query(None),
    _: Any = Depends(require_role("viewer")),
) -> SimulationResponse:
    return SimulationResponse(**_build_score_simulation(
        gstin,
        company_name,
        request_id=getattr(request.state, "request_id", None),
    ))


@router.post(
    "/score/{gstin}/refresh",
    summary="Force refresh pipeline data for a GSTIN",
)
@router.post(
    "/v1/score/{gstin}/refresh",
    summary="Force refresh pipeline data for a GSTIN (v1)",
)
async def refresh_score_inputs(
    gstin: str,
    company_name: Optional[str] = Query(None),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    normalized_gstin = normalize_gstin(gstin)
    if not is_valid_gstin(normalized_gstin):
        raise HTTPException(status_code=400, detail="Invalid GSTIN format")

    refresh_gstin_now(normalized_gstin, company_name)
    pipeline_data = get_storage().get_pipeline_data(normalized_gstin)
    return {
        "status": "refreshed",
        "gstin": normalized_gstin,
        "pipeline_timestamp": pipeline_data.get("pipeline_timestamp") if pipeline_data else None,
    }


@router.post(
    "/score/{gstin}/refresh/{stream}",
    summary="Force refresh a single pipeline stream for a GSTIN",
)
@router.post(
    "/v1/score/{gstin}/refresh/{stream}",
    summary="Force refresh a single pipeline stream for a GSTIN (v1)",
)
async def refresh_score_stream(
    gstin: str,
    stream: str,
    company_name: Optional[str] = Query(None),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    normalized_gstin = normalize_gstin(gstin)
    if not is_valid_gstin(normalized_gstin):
        raise HTTPException(status_code=400, detail="Invalid GSTIN format")

    normalized_stream = stream.strip().lower()
    if normalized_stream not in {"gst", "upi", "eway"}:
        raise HTTPException(status_code=400, detail="Invalid stream. Must be one of: gst, upi, eway")

    canonical = refresh_pipeline_stream(normalized_gstin, normalized_stream, company_name)
    pipeline_data = get_storage().get_pipeline_data(normalized_gstin)
    stream_payload = (pipeline_data or {}).get(canonical, {})
    return {
        "status": "refreshed",
        "gstin": normalized_gstin,
        "stream": normalized_stream,
        "pipeline_type": canonical,
        "fetched_at": stream_payload.get("data_freshness"),
    }


@router.get(
    "/score/{gstin}/percentile",
    summary="Return score and feature percentiles for a GSTIN",
)
@router.get(
    "/v1/score/{gstin}/percentile",
    summary="Return score and feature percentiles for a GSTIN (v1)",
)
async def get_score_percentile(
    gstin: str,
    request: Request,
    company_name: Optional[str] = Query(None),
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    payload = _score_assessment_payload(
        gstin,
        company_name,
        request_id=getattr(request.state, "request_id", None),
        persist=False,
    )

    return {
        "gstin": payload["gstin"],
        **payload["percentile"],
        "feature_percentiles": {},
    }


@router.get(
    "/score/{gstin}/export.docx",
    summary="Export a DOCX scorecard for a GSTIN",
)
@router.get(
    "/v1/score/{gstin}/export.docx",
    summary="Export a DOCX scorecard for a GSTIN (v1)",
)
async def export_score_docx(
    gstin: str,
    request: Request,
    company_name: Optional[str] = Query(None),
    industry_code: Optional[str] = Query(None),
    _: Any = Depends(require_role("viewer")),
) -> StreamingResponse:
    payload = _score_assessment_payload(
        gstin,
        company_name,
        industry_code,
        request_id=getattr(request.state, "request_id", None),
        persist=False,
    )
    stream = _render_docx(payload)
    filename = f"MSME_Scorecard_{payload['gstin']}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post(
    "/model/retrain",
    summary="Retrain and persist the MSME scoring model",
    response_model=ModelRetrainResponse,
)
@router.post(
    "/v1/model/retrain",
    summary="Retrain and persist the MSME scoring model (v1)",
    response_model=ModelRetrainResponse,
)
async def retrain_model(
    body: ModelRetrainRequest,
    _: Any = Depends(require_role("admin")),
) -> ModelRetrainResponse:
    storage = get_storage()
    recorded_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    outcomes_recorded = 0
    for outcome in body.outcomes:
        normalized_gstin = normalize_gstin(outcome.gstin)
        if not is_valid_gstin(normalized_gstin):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid GSTIN format in outcome payload: {outcome.gstin}",
            )
        feature_snapshot = _build_outcome_feature_snapshot(normalized_gstin)
        storage.record_loan_outcome(
            gstin=normalized_gstin,
            company_name=outcome.company_name,
            repaid=outcome.outcome == "repaid",
            loan_amount=float(outcome.loan_amount),
            tenure_months=int(outcome.tenure_months),
            recorded_at=recorded_at,
            source_model_version=storage.get_latest_assessment_model_version(normalized_gstin),
            feature_snapshot=feature_snapshot,
        )
        outcomes_recorded += 1

    total_real_outcomes = storage.count_loan_outcomes()
    scorer = get_scorer()
    previous_health = scorer.health_summary()
    previous_auc = previous_health.get("metrics", {}).get("auc")
    real_outcomes = storage.get_loan_outcomes()

    if total_real_outcomes < REAL_OUTCOME_MIN_RECORDS:
        return ModelRetrainResponse(
            status="outcomes_recorded_pending_more_feedback",
            model=previous_health,
            retrained=False,
            outcomes_recorded=outcomes_recorded,
            total_real_outcomes=total_real_outcomes,
            governance={
                "minimum_real_outcomes_required": REAL_OUTCOME_MIN_RECORDS,
                "real_outcomes_used": 0,
                "synthetic_outcomes_used": previous_health.get("metrics", {})
                .get("synthetic_dataset", {})
                .get("panel_rows", 0),
                "auc_before": previous_auc,
                "auc_after": previous_auc,
                "feature_schema_version": previous_health.get("metrics", {}).get("feature_schema_version"),
            },
        )

    retrain_result = scorer.retrain(real_outcomes=real_outcomes)
    training_summary = retrain_result.get("training_summary", {})
    auc_after = retrain_result.get("metrics", {}).get("auc")
    governance = {
        "minimum_real_outcomes_required": REAL_OUTCOME_MIN_RECORDS,
        "real_outcomes_used": training_summary.get("real_sample_count", 0),
        "synthetic_outcomes_used": training_summary.get("synthetic_sample_count", 0),
        "training_sample_size": training_summary.get("training_sample_size", 0),
        "real_label_ratio": training_summary.get("real_label_ratio", 0.0),
        "auc_before": previous_auc,
        "auc_after": auc_after,
        "feature_schema_version": training_summary.get("feature_schema_version"),
    }
    storage.record_model_version(
        model_version=retrain_result["model_version"],
        trained_at=retrain_result["last_loaded_at"],
        training_sample_size=training_summary.get("training_sample_size", 0),
        synthetic_sample_count=training_summary.get("synthetic_sample_count", 0),
        real_sample_count=training_summary.get("real_sample_count", 0),
        real_label_ratio=training_summary.get("real_label_ratio", 0.0),
        auc_before=previous_auc,
        auc_after=auc_after,
        feature_schema_version=training_summary.get("feature_schema_version", "unknown"),
        metrics=retrain_result.get("metrics", {}),
    )

    return ModelRetrainResponse(
        status="retrained",
        model=retrain_result,
        retrained=True,
        outcomes_recorded=outcomes_recorded,
        total_real_outcomes=total_real_outcomes,
        governance=governance,
    )


# ──────────────────────────────────────────────────────────
#  ENTITY GRAPH — visualizes the fraud network topology
# ──────────────────────────────────────────────────────────

def _resolve_directors(gstin: str) -> list[Dict[str, Any]]:
    """Resolve director data from MCA fixtures for demo GSTINs."""
    from ..fixtures.agent_fixtures import MCA_FIXTURES

    demo_map = {
        "27ARJUN1234A1Z5": "reject",
        "29CLEAN5678B1Z2": "approve",
    }
    scenario = demo_map.get(gstin)
    if scenario and scenario in MCA_FIXTURES:
        return MCA_FIXTURES[scenario].get("directors", [])
    return []


def _resolve_company_name(gstin: str) -> str:
    demo_names = {
        "27ARJUN1234A1Z5": "Arjun Textiles Pvt. Ltd.",
        "29CLEAN5678B1Z2": "CleanTech Manufacturing Ltd.",
        "09NEWCO1234A1Z9": "New Startup Pvt. Ltd.",
    }
    return demo_names.get(gstin, f"MSME Entity ({gstin[:5]}...)")


@router.get(
    "/entity-graph/{gstin}",
    summary="Entity relationship graph for a GSTIN",
)
@router.get(
    "/v1/entity-graph/{gstin}",
    summary="Entity relationship graph for a GSTIN (v1)",
    description="Returns nodes (companies, directors, counterparties) and edges (UPI flows, director links) for fraud network visualization.",
)
async def get_entity_graph(
    gstin: str,
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    normalized_gstin = normalize_gstin(gstin)
    if not is_valid_gstin(normalized_gstin):
        raise HTTPException(status_code=400, detail="Invalid GSTIN format")

    storage = get_storage()
    pipeline_data = storage.get_pipeline_data(normalized_gstin)
    if pipeline_data is None:
        trigger_immediate_ingestion(normalized_gstin)
        retry_after = min(get_settings().pipeline_interval_seconds, 30)
        raise HTTPException(
            status_code=202,
            detail="Entity graph is not ready yet because pipeline ingestion is still in progress for this GSTIN.",
            headers={"Retry-After": str(retry_after)},
        )

    upi_transactions = pipeline_data.get("upi_cadence", {}).get("transactions", [])

    detector = UPIFraudDetector()
    company_name = _resolve_company_name(normalized_gstin)
    directors = _resolve_directors(normalized_gstin)
    graph, cycles = detector.build_networkx_entity_graph(
        transactions=upi_transactions,
        gstin=normalized_gstin,
        company_name=company_name,
        directors=directors,
        entity_graph_service=get_entity_graph_service(),
    )
    return serialize_graph(
        G=graph,
        queried_gstin=normalized_gstin,
        detected_cycles=cycles,
    )


@router.get(
    "/score/{gstin}/fraud-graph",
    summary="Cross-GSTIN fraud graph for a GSTIN",
)
@router.get(
    "/v1/score/{gstin}/fraud-graph",
    summary="Cross-GSTIN fraud graph for a GSTIN (v1)",
)
async def get_fraud_graph(
    gstin: str,
    _: Any = Depends(require_role("viewer")),
) -> Dict[str, Any]:
    graph = await get_entity_graph(gstin, _)
    normalized_gstin = normalize_gstin(gstin)

    node_entries = graph.get("nodes", [])
    edge_entries = graph.get("edges", [])
    cycle_paths = graph.get("meta", {}).get("cycle_paths", [])
    cycle_label_set = {node for cycle in cycle_paths for node in cycle}

    nodes = [
        {
            "id": node["data"]["label"],
            "is_queried": node["data"].get("is_queried", False),
            "in_cycle": node["data"].get("in_cycle", False),
            "entity_type": node["data"].get("type", "business"),
        }
        for node in node_entries
        if node["data"].get("type") != "director"
    ]
    edges = [
        {
            "source": next((node["data"]["label"] for node in node_entries if node["data"]["id"] == edge["data"]["source"]), edge["data"]["source"]),
            "target": next((node["data"]["label"] for node in node_entries if node["data"]["id"] == edge["data"]["target"]), edge["data"]["target"]),
            "amount": edge["data"].get("amount", 0),
            "direction": "outflow" if edge["data"]["source"] == normalized_gstin else "inflow",
            "in_cycle": edge["data"].get("in_cycle", False),
        }
        for edge in edge_entries
        if edge["data"].get("type") != "director_of"
    ]

    return {
        "gstin": normalized_gstin,
        "nodes": nodes,
        "edges": edges,
        "cycles": cycle_paths,
        "fraud_detected": graph.get("meta", {}).get("cycles_detected", 0) > 0,
        "risk_summary": f"Detected {graph.get('meta', {}).get('cycles_detected', 0)} circular topology path(s)." if cycle_label_set else "No circular topology detected.",
        "graph_stats": {
            "nodes": graph.get("meta", {}).get("total_nodes", 0),
            "edges": graph.get("meta", {}).get("total_edges", 0),
            "density": 0,
        },
    }
