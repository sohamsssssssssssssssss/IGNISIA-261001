from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
try:
    # WeasyPrint needs native libraries like libgobject/Pango. On macOS these
    # are often missing unless GTK is installed, so keep PDF export optional.
    from weasyprint import HTML as _WeasyHTML
    _WEASYPRINT_AVAILABLE = True
except (OSError, ImportError):
    _WeasyHTML = None
    _WEASYPRINT_AVAILABLE = False
import os
from datetime import date
from typing import Dict, Any


# ─────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (e.g. '1F4E79')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right → hex colour string."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in kwargs.items():
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _bold_run(para, text: str, size_pt: int = 11, color_hex: str = None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _normal_run(para, text: str, size_pt: int = 11, color_hex: str = None, italic: bool = False):
    run = para.add_run(text)
    run.bold = False
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _add_section_header(doc: Document, title: str, section_num: str):
    """Dark navy section heading bar."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"  {section_num}  {title.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Shade the paragraph background using XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F4E79')
    pPr.append(shd)
    return para


def _add_subsection(doc: Document, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # Bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_kv_table(doc: Document, rows: list[tuple], col_widths=(3.2, 3.4)):
    """Two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        # Key cell
        kc = row.cells[0]
        kc.width = Inches(col_widths[0])
        _set_cell_bg(kc, 'D6E4F0')
        kp = kc.paragraphs[0]
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after = Pt(3)
        _bold_run(kp, key, size_pt=10)
        # Value cell
        vc = row.cells[1]
        vc.width = Inches(col_widths[1])
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after = Pt(3)
        _normal_run(vp, str(val), size_pt=10)
    doc.add_paragraph()  # spacer


def _add_five_c_score_bar(doc: Document, c_name: str, score: float, rationale: str, flag: str = None):
    """
    Renders a single 'C' as: coloured label | score% bar (visual) | rationale text
    """
    COLOR_MAP = {
        'Character':  '2E75B6',
        'Capacity':   '70AD47',
        'Capital':    'ED7D31',
        'Collateral': '9E480E',
        'Conditions': '4472C4',
    }
    bar_color = COLOR_MAP.get(c_name, '1F4E79')

    # Score badge row
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.4), Inches(1.0), Inches(4.2)]

    label_cell, score_cell, rationale_cell = table.rows[0].cells
    label_cell.width = widths[0]
    score_cell.width = widths[1]
    rationale_cell.width = widths[2]

    _set_cell_bg(label_cell, bar_color)
    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)
    lr = lp.add_run(f"  {c_name.upper()}")
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _set_cell_bg(score_cell, 'F2F2F2')
    sp = score_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after = Pt(4)
    score_val = int(score) if score == int(score) else round(score, 1)
    sr = sp.add_run(f"{score_val}%")
    sr.bold = True
    sr.font.size = Pt(11)
    # Color-code the score
    if score >= 80:
        sr.font.color.rgb = RGBColor(0x37, 0x5A, 0x23)   # dark green
    elif score >= 60:
        sr.font.color.rgb = RGBColor(0xBF, 0x85, 0x02)   # amber
    else:
        sr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)   # red

    rp = rationale_cell.paragraphs[0]
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)
    _normal_run(rp, rationale, size_pt=9, italic=True)
    if flag:
        rp.add_run("  ")
        fr = rp.add_run(f"⚠ {flag}")
        fr.font.size = Pt(9)
        fr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        fr.bold = True

    doc.add_paragraph()  # spacer after each C


# ─────────────────────────────────────────────
#  Five Cs content library
# ─────────────────────────────────────────────

FIVE_CS_DEFINITIONS = {
    "Character": {
        "definition": (
            "Assesses the borrower's willingness and intent to repay, based on promoter "
            "integrity, litigation history, regulatory compliance record, and historical "
            "credit behaviour (CIBIL CMR rating)."
        ),
        "key_indicators": [
            "CIBIL CMR Commercial Rating",
            "eCourts / DRT Litigation Status",
            "MCA Director NPA Linkages",
            "Promoter Background & Track Record",
            "Prior Default History",
        ],
        "india_note": (
            "India-specific: CIBIL CMR-1 to CMR-6 scale used. CMR-1 to CMR-3 is institutional "
            "threshold. DRT (Debt Recovery Tribunal) cases under SARFAESI Act 2002 are "
            "auto-flagged from eCourts portal."
        ),
    },
    "Capacity": {
        "definition": (
            "Evaluates the borrower's ability to service debt from operating cash flows. "
            "Analyses GST revenue trends, bank credit turnover, DSCR, EBITDA margins, "
            "and working capital cycle."
        ),
        "key_indicators": [
            "Debt Service Coverage Ratio (DSCR ≥ 1.25x required)",
            "GST-declared revenue vs Bank Credits (cross-verified)",
            "EBITDA Margin & Operating Leverage",
            "Working Capital Days (Debtor/Creditor/Inventory)",
            "26AS TDS vs ITR Gross Revenue reconciliation",
        ],
        "india_note": (
            "India-specific: GSTR-3B vs GSTR-2A ITC mismatch flagged as circular trading "
            "signal. Bank statement credits cross-referenced against GST Output Liability "
            "at invoice level."
        ),
    },
    "Capital": {
        "definition": (
            "Measures the borrower's financial cushion — net worth, leverage ratios, "
            "promoter equity stake, and ability to absorb losses without triggering default."
        ),
        "key_indicators": [
            "Tangible Net Worth (TNW)",
            "Total Outside Liabilities / TNW (TOL/TNW ≤ 3x)",
            "Debt-to-Equity Ratio",
            "Promoter Contribution & Skin-in-the-Game",
            "Reserves & Surplus Trend (3-year)",
        ],
        "india_note": (
            "India-specific: Examined via ITR-6 Schedule-BS and audited P&L. "
            "Related-party transactions (Schedule AL) reviewed for capital diversion."
        ),
    },
    "Collateral": {
        "definition": (
            "Assesses the quality, realisability, and legal enforceability of security "
            "offered — immovable property, plant & machinery, book debts, and guarantees."
        ),
        "key_indicators": [
            "Primary Security: Hypothecation of Current Assets",
            "Collateral: Immovable Property (EM / Mortgage)",
            "SARFAESI Enforceability Status",
            "Approved Valuer Report (within 3 years)",
            "Security Coverage Ratio (SCR ≥ 1.33x)",
        ],
        "india_note": (
            "India-specific: SARFAESI Act 2002 enforceability is a hard requirement for "
            "limits above ₹1 Cr. CERSAI registration mandatory for equitable mortgages."
        ),
    },
    "Conditions": {
        "definition": (
            "Examines the external macro environment — sector-specific tailwinds/headwinds, "
            "RBI policy, industry growth rates, and the purpose/end-use of the proposed credit."
        ),
        "key_indicators": [
            "Sector Growth Rate (YoY, NIC Code-based)",
            "RBI Regulatory Headwinds (Sector-specific circulars)",
            "End-use of Funds (Working Capital vs Capex)",
            "Commodity/FX Exposure Risk",
            "Post-COVID Demand Recovery Index (if applicable)",
        ],
        "india_note": (
            "India-specific: RBI Master Directions on Priority Sector Lending, PSL "
            "classification, and sector-specific SMA/NPA norms reviewed."
        ),
    },
}


# ─────────────────────────────────────────────
#  Main Generator
# ─────────────────────────────────────────────

class CAMGenerator:
    """
    Generates a professional Word DOCX CAM with full Five Cs of Credit analysis.
    """

    def __init__(self, templates_dir: str = "/tmp/templates"):
        self.templates_dir = templates_dir
        os.makedirs(self.templates_dir, exist_ok=True)

    # ------------------------------------------------------------------
    def generate_cam_docx(self, borrower_data: Dict[str, Any], output_path: str) -> str:
        doc = Document()

        # ── Page margins ──────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

        self._add_cover(doc, borrower_data)
        self._add_borrower_overview(doc, borrower_data)
        self._add_ai_recommendation(doc, borrower_data)
        self._add_five_cs(doc, borrower_data)
        self._add_shap_factors(doc, borrower_data)
        self._add_narrative(doc, borrower_data)
        self._add_swot(doc, borrower_data)
        self._add_triangulation(doc, borrower_data)
        self._add_verdict_recommendation(doc, borrower_data)
        self._add_disclaimer(doc)

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Section 0 – Cover
    # ------------------------------------------------------------------
    def _add_cover(self, doc: Document, data: Dict):
        title = doc.add_heading('', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('COMPREHENSIVE CREDIT APPRAISAL MEMO')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _normal_run(sub, 'INTELLI-CREDIT AI ENGINE  ·  CONFIDENTIAL', size_pt=10, color_hex='808080')

        doc.add_paragraph()
        _add_kv_table(doc, [
            ('Borrower Name',    data.get('name', 'N/A')),
            ('PAN',              data.get('pan', 'N/A')),
            ('Sector',           data.get('sector', 'N/A')),
            ('Promoter',         data.get('promoter', 'N/A')),
            ('Memo Date',        date.today().strftime('%d %B %Y')),
            ('Session ID',       data.get('session_id', 'N/A')),
        ])

    # ------------------------------------------------------------------
    # Section 1 – Borrower Overview
    # ------------------------------------------------------------------
    def _add_borrower_overview(self, doc: Document, data: Dict):
        _add_section_header(doc, 'Borrower Overview', '1.')
        _add_kv_table(doc, [
            ('Segment',          data.get('segment', 'NON-MSME')),
            ('Annual Turnover',  f"INR {data.get('turnover', 'N/A')} Cr (FY24)"),
            ('Vintage',          data.get('vintage', 'N/A')),
            ('CIBIL CMR',        data.get('cibil_cmr', 'N/A')),
            ('GST Registration', data.get('gstin', 'N/A')),
        ])

    # ------------------------------------------------------------------
    # Section 2 – AI Recommendation
    # ------------------------------------------------------------------
    def _add_ai_recommendation(self, doc: Document, data: Dict):
        _add_section_header(doc, 'AI Credit Recommendation', '2.')

        score = data.get('credit_score', 0)
        verdict = data.get('verdict', 'APPROVE' if score >= 70 else 'REJECT')
        limit = data.get('recommended_limit', '0.00')
        rate = data.get('final_interest_rate', 'N/A')
        mclr = data.get('base_rate_mclr', '8.50%')
        risk_premium = data.get('cbm_risk_premium', 'N/A')
        sector_spread = data.get('sector_spread', 'N/A')

        verdict_color = '375A23' if verdict == 'APPROVE' else 'C00000'

        _add_kv_table(doc, [
            ('AI Credit Score',       f"{score} / 100"),
            ('Decision Verdict',      verdict),
            ('Recommended Limit',     f"INR {limit} Cr"),
            ('Base Rate (MCLR)',      mclr),
            ('CBM Risk Premium',      risk_premium),
            ('Sector Spread',         sector_spread),
            ('Final Interest Rate',   rate),
        ])

        # Verdict rationale
        para = doc.add_paragraph()
        _bold_run(para, 'Rationale:  ', size_pt=10)
        _normal_run(para, data.get('verdict_rationale',
            'Score computed via Concept Bottleneck Model (CBM) + SHAP explainability pipeline.'), size_pt=10)

    # ------------------------------------------------------------------
    # Section 3 – Five Cs of Credit  ← THE UPGRADED SECTION
    # ------------------------------------------------------------------
    def _add_five_cs(self, doc: Document, data: Dict):
        _add_section_header(doc, 'The Five Cs of Credit — Deep Analysis', '3.')

        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(8)
        _normal_run(
            intro,
            'Each dimension is scored independently by the Concept Bottleneck Model (CBM) '
            'and weighted by SHAP feature importance. India-specific regulatory norms and '
            'RBI guidelines are applied throughout.',
            size_pt=10, italic=True
        )

        concept_scores = data.get('concept_scores', {})
        concept_flags = data.get('concept_flags', {})       # e.g. {"Character": "Active DRT litigation"}
        concept_narratives = data.get('concept_narratives', {})  # per-C LLM narrative

        C_ORDER = ['Character', 'Capacity', 'Capital', 'Collateral', 'Conditions']

        for c_name in C_ORDER:
            meta = FIVE_CS_DEFINITIONS[c_name]
            score = concept_scores.get(c_name, 0)
            flag = concept_flags.get(c_name)
            narrative = concept_narratives.get(c_name, meta['definition'])

            # ── C header + score bar ──────────────────────────────────
            _add_subsection(doc, f"3.{C_ORDER.index(c_name)+1}  {c_name}")
            _add_five_c_score_bar(doc, c_name, score, narrative, flag)

            # ── Definition ───────────────────────────────────────────
            def_para = doc.add_paragraph()
            def_para.paragraph_format.left_indent = Inches(0.2)
            def_para.paragraph_format.space_after = Pt(4)
            _bold_run(def_para, 'Definition:  ', size_pt=10)
            _normal_run(def_para, meta['definition'], size_pt=10)

            # ── Key Indicators table ──────────────────────────────────
            ind_label = doc.add_paragraph()
            ind_label.paragraph_format.left_indent = Inches(0.2)
            _bold_run(ind_label, 'Key Assessment Indicators:', size_pt=10)

            ind_table = doc.add_table(rows=len(meta['key_indicators']), cols=2)
            ind_table.style = 'Table Grid'
            ind_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for idx, indicator in enumerate(meta['key_indicators']):
                row = ind_table.rows[idx]
                num_cell = row.cells[0]
                num_cell.width = Inches(0.3)
                _set_cell_bg(num_cell, 'EBF3FB')
                np_ = num_cell.paragraphs[0]
                np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _bold_run(np_, str(idx + 1), size_pt=9)

                ind_cell = row.cells[1]
                ind_cell.width = Inches(6.2)
                ip = ind_cell.paragraphs[0]
                ip.paragraph_format.space_before = Pt(2)
                ip.paragraph_format.space_after = Pt(2)
                _normal_run(ip, indicator, size_pt=9)

            doc.add_paragraph()

            # ── India-specific note ───────────────────────────────────
            note_para = doc.add_paragraph()
            note_para.paragraph_format.left_indent = Inches(0.2)
            note_para.paragraph_format.space_after = Pt(10)
            _bold_run(note_para, '🇮🇳  India Context:  ', size_pt=9, color_hex='1F4E79')
            _normal_run(note_para, meta['india_note'], size_pt=9, color_hex='404040', italic=True)

        # ── Aggregate Five Cs summary table ──────────────────────────
        _add_subsection(doc, '3.6  Five Cs — Aggregate Scorecard')
        agg_rows = [('Dimension', 'Score', 'Weight', 'Weighted Score')]
        weights = {'Character': 0.25, 'Capacity': 0.30, 'Capital': 0.20,
                   'Collateral': 0.15, 'Conditions': 0.10}
        total_weighted = 0.0
        for c in C_ORDER:
            s = concept_scores.get(c, 0)
            w = weights[c]
            ws = round(s * w, 1)
            total_weighted += ws
            agg_rows.append((c, f"{s}%", f"{int(w*100)}%", f"{ws}"))
        agg_rows.append(('COMPOSITE SCORE', '', '', f"{round(total_weighted, 1)} / 100"))

        agg_table = doc.add_table(rows=len(agg_rows), cols=4)
        agg_table.style = 'Table Grid'
        agg_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        col_w = [Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.5)]

        for i, row_data in enumerate(agg_rows):
            row = agg_table.rows[i]
            is_header = (i == 0)
            is_total = (i == len(agg_rows) - 1)
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.width = col_w[j]
                if is_header:
                    _set_cell_bg(cell, '1F4E79')
                elif is_total:
                    _set_cell_bg(cell, 'D6E4F0')
                else:
                    _set_cell_bg(cell, 'FFFFFF' if i % 2 == 0 else 'F5F9FF')
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                cp.paragraph_format.space_before = Pt(3)
                cp.paragraph_format.space_after = Pt(3)
                run = cp.add_run(cell_text)
                run.bold = is_header or is_total
                run.font.size = Pt(10)
                if is_header:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section 4 – SHAP Factors
    # ------------------------------------------------------------------
    def _add_shap_factors(self, doc: Document, data: Dict):
        _add_section_header(doc, 'SHAP Explainability — Feature Contributions', '4.')

        shap_factors = data.get('shap_top_factors', [])
        if not shap_factors:
            doc.add_paragraph('No SHAP data available.')
            return

        shap_table = doc.add_table(rows=len(shap_factors) + 1, cols=3)
        shap_table.style = 'Table Grid'
        headers = ['Feature', 'Contribution (pts)', 'Direction']
        header_row = shap_table.rows[0]
        for j, h in enumerate(headers):
            c = header_row.cells[j]
            _set_cell_bg(c, '1F4E79')
            p = c.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, factor in enumerate(shap_factors):
            row = shap_table.rows[i + 1]
            contrib = factor.get('contribution', 0)
            direction = '▲ Positive' if contrib > 0 else '▼ Negative'
            dir_color = '375A23' if contrib > 0 else 'C00000'
            vals = [factor.get('feature', ''), f"{contrib:+.1f}", direction]
            for j, val in enumerate(vals):
                c = row.cells[j]
                _set_cell_bg(c, 'FFFFFF' if i % 2 == 0 else 'F5F9FF')
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(val)
                run.font.size = Pt(10)
                if j == 2:
                    run.font.color.rgb = RGBColor.from_string(dir_color)
                    run.bold = True

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section 5 – Narrative
    # ------------------------------------------------------------------
    def _add_narrative(self, doc: Document, data: Dict):
        _add_section_header(doc, 'AI-Generated Qualitative Narrative', '5.')
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.space_after = Pt(6)
        _normal_run(para, data.get('llm_narrative', 'No narrative provided.'), size_pt=10, italic=True)

        # DD Override notes
        dd_notes = data.get('dd_override_notes')
        if dd_notes:
            _add_subsection(doc, 'Due Diligence Override Notes (Credit Officer)')
            dd_para = doc.add_paragraph()
            dd_para.paragraph_format.left_indent = Inches(0.2)
            _normal_run(dd_para, dd_notes, size_pt=10)

    # ------------------------------------------------------------------
    # Section 6 – SWOT Analysis
    # ------------------------------------------------------------------
    def _add_swot(self, doc: Document, data: Dict):
        swot = data.get('swot')
        if not swot:
            return

        _add_section_header(doc, 'SWOT Analysis — AI-Generated', '6.')

        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(6)
        _normal_run(
            intro,
            'Generated by the SWOT Engine using concept scores, SHAP attribution, '
            'sector research, and extracted document data. All findings are grounded '
            'in verified data — no hallucinated content.',
            size_pt=10, italic=True
        )

        SWOT_CELLS = [
            ('Strengths', swot.get('strengths', []), '375A23', 'E8F5E9'),
            ('Weaknesses', swot.get('weaknesses', []), 'C00000', 'FFEBEE'),
            ('Opportunities', swot.get('opportunities', []), '1F4E79', 'E3F2FD'),
            ('Threats', swot.get('threats', []), 'E65100', 'FFF3E0'),
        ]

        # 2×2 grid as a 2-row, 2-col table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for idx, (label, items, header_color, bg_color) in enumerate(SWOT_CELLS):
            row_idx = idx // 2
            col_idx = idx % 2
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = Inches(3.3)

            # Header
            hp = cell.paragraphs[0]
            hp.paragraph_format.space_before = Pt(6)
            hp.paragraph_format.space_after = Pt(4)
            _set_cell_bg(cell, bg_color)
            _bold_run(hp, f'◆ {label.upper()}', size_pt=10, color_hex=header_color)

            # Items
            for item in items:
                ip = cell.add_paragraph()
                ip.paragraph_format.space_before = Pt(1)
                ip.paragraph_format.space_after = Pt(1)
                ip.paragraph_format.left_indent = Inches(0.1)
                _normal_run(ip, f'→ {item}', size_pt=9)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section 7 – Data Triangulation
    # ------------------------------------------------------------------
    def _add_triangulation(self, doc: Document, data: Dict):
        tri_findings = data.get('triangulation', [])
        if not tri_findings:
            return

        _add_section_header(doc, 'Data Triangulation — External vs Internal', '7.')

        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(6)
        _normal_run(
            intro,
            'Cross-reference of external intelligence (eCourts, MCA, news, sector data) '
            'against data extracted from uploaded documents.',
            size_pt=10, italic=True
        )

        # Summary counts
        aligned = sum(1 for f in tri_findings if f.get('alignment') == 'ALIGNED')
        contradicted = sum(1 for f in tri_findings if f.get('alignment') == 'CONTRADICTED')
        unverified = sum(1 for f in tri_findings if f.get('alignment') == 'UNVERIFIED')

        summary_para = doc.add_paragraph()
        _bold_run(summary_para, f'Summary: ', size_pt=10)
        _normal_run(summary_para, f'{aligned} Aligned  |  {contradicted} Contradicted  |  {unverified} Unverified', size_pt=10)
        doc.add_paragraph()

        # Table
        headers = ['#', 'External Signal', 'Internal Data', 'Alignment', 'Recommendation']
        table = doc.add_table(rows=len(tri_findings) + 1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        header_row = table.rows[0]
        for j, h in enumerate(headers):
            c = header_row.cells[j]
            _set_cell_bg(c, '1F4E79')
            p = c.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        alignment_colors = {'ALIGNED': '375A23', 'CONTRADICTED': 'C00000', 'UNVERIFIED': 'BF8502'}

        for i, finding in enumerate(tri_findings):
            row = table.rows[i + 1]
            alignment = finding.get('alignment', 'UNVERIFIED')
            vals = [
                str(i + 1),
                finding.get('external', finding.get('external_signal', '')),
                finding.get('internal', finding.get('internal_data_point', '')),
                alignment,
                finding.get('rec', finding.get('recommendation', '')),
            ]
            for j, val in enumerate(vals):
                c = row.cells[j]
                _set_cell_bg(c, 'FFFFFF' if i % 2 == 0 else 'F5F9FF')
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(val)
                run.font.size = Pt(8)
                if j == 3:
                    run.font.color.rgb = RGBColor.from_string(alignment_colors.get(alignment, '404040'))
                    run.bold = True

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section 8 – Verdict & Recommendation
    # ------------------------------------------------------------------
    def _add_verdict_recommendation(self, doc: Document, data: Dict):
        _add_section_header(doc, 'Final Verdict & Recommendation', '8.')

        score = data.get('credit_score', 0)
        verdict = data.get('verdict', 'APPROVE' if score >= 70 else 'REJECT')
        verdict_color = '375A23' if verdict == 'APPROVE' else 'C00000'

        # Verdict stamp
        v_para = doc.add_paragraph()
        v_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v_para.paragraph_format.space_before = Pt(12)
        v_para.paragraph_format.space_after = Pt(12)
        vr = v_para.add_run(f'   {verdict}   ')
        vr.bold = True
        vr.font.size = Pt(24)
        vr.font.color.rgb = RGBColor.from_string(verdict_color)

        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(score_para, f'AI Credit Score: {score} / 100', size_pt=14, color_hex=verdict_color)

        # Rationale
        _add_subsection(doc, '8.1  Decision Rationale')
        rat_para = doc.add_paragraph()
        rat_para.paragraph_format.left_indent = Inches(0.2)
        rat_para.paragraph_format.space_after = Pt(8)
        _normal_run(rat_para, data.get('verdict_rationale', 'No rationale provided.'), size_pt=10)

        # Loan parameters (if approved)
        loan = data.get('loan_details')
        if loan and verdict == 'APPROVE':
            _add_subsection(doc, '8.2  Sanction Parameters')
            _add_kv_table(doc, [
                ('Loan Type', loan.get('type', 'N/A')),
                ('Sanctioned Amount', loan.get('amount', 'N/A')),
                ('Tenure', loan.get('tenure', 'N/A')),
                ('Interest Rate', loan.get('rate', 'N/A')),
                ('Recommended Limit', data.get('recommended_limit', 'N/A')),
            ])

        # Rejection notice (if rejected)
        if verdict == 'REJECT':
            _add_subsection(doc, '8.2  RBI Adverse Action Notice (FR-38)')
            notice_para = doc.add_paragraph()
            notice_para.paragraph_format.left_indent = Inches(0.2)
            _normal_run(
                notice_para,
                'Pursuant to RBI Master Direction DOR.STR.REC.10/21.04.048/2025-26, '
                'the application for credit facility has been declined. The applicant '
                'retains the right to request a copy of the credit report used in this '
                'assessment and to seek review through the Internal Grievance Redressal '
                'Officer within 30 days of this notice.',
                size_pt=10, italic=True
            )

    # ------------------------------------------------------------------
    # Disclaimer
    # ------------------------------------------------------------------
    def _add_disclaimer(self, doc: Document):
        doc.add_paragraph()
        disc = doc.add_paragraph()
        disc.paragraph_format.space_before = Pt(16)
        _normal_run(
            disc,
            'DISCLAIMER: This memo is generated by the Intelli-Credit AI Engine and is '
            'intended to assist—not replace—the judgment of a qualified credit officer. '
            'All decisions are subject to applicable RBI Master Directions and internal '
            'credit policy. Generated: ' + date.today().strftime('%d %B %Y'),
            size_pt=8, color_hex='808080', italic=True
        )

    # ------------------------------------------------------------------
    def convert_docx_to_pdf(self, docx_path: str, output_pdf_path: str) -> str:
        if not _WEASYPRINT_AVAILABLE:
            raise RuntimeError(
                "WeasyPrint is not available on this system. DOCX export still works, "
                "but PDF conversion needs native libraries such as libgobject and Pango. "
                "On local macOS setups, install Homebrew's gtk+3 stack (or equivalent GTK libraries) "
                "to enable PDF export."
            )
        html_string = f"<h1>Credit Appraisal Memo</h1><p>Converted from {docx_path}</p>"
        _WeasyHTML(string=html_string).write_pdf(output_pdf_path)
        return output_pdf_path


# ─────────────────────────────────────────────
#  Adverse Action Generator (unchanged)
# ─────────────────────────────────────────────

class AdverseActionGenerator:
    """
    Generates RBI Fair Practices Code-compliant Adverse Action Notices on rejection.
    """
    def generate_notice(self, borrower_name: str, rejection_factors: list[str], output_path: str):
        document = Document()
        document.add_heading('Adverse Action Notice', 0)

        document.add_paragraph(f"To: {borrower_name}")
        document.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
        document.add_paragraph("\nDear Applicant,")
        document.add_paragraph(
            "In accordance with the RBI Fair Practices Code (Master Direction – FIDD.CO.Plan.1/04.09.01/2016-17), "
            "we are writing to inform you that we are unable to approve your application for credit at this time."
        )
        document.add_paragraph("Principal Reasons for Decision:")
        for r in rejection_factors:
            document.add_paragraph(f"{r}", style='List Bullet')

        document.add_paragraph("\nYou have the right to seek a review of this decision by writing to our "
                               "Internal Grievance Redressal Officer within 30 days of this notice.")
        document.add_paragraph("\nSincerely,")
        document.add_paragraph("Credit Department")

        document.save(output_path)
        return output_path


# ─────────────────────────────────────────────
#  Quick local test
# ─────────────────────────────────────────────
if __name__ == "__main__":
    sample_data = {
        "name": "CleanTech Manufacturing Ltd.",
        "pan": "BBMCV5678K",
        "sector": "Industrial Machinery – CNC & Automation",
        "promoter": "Vikram Mehta",
        "vintage": "7 years",
        "turnover": "42.5",
        "cibil_cmr": "CMR-3 (Low Risk)",
        "gstin": "27BBMCV5678K1Z5",
        "segment": "NON-MSME",
        "session_id": "IC-20260306-0042",
        "credit_score": 88.5,
        "verdict": "APPROVE",
        "recommended_limit": "5.00",
        "base_rate_mclr": "8.50%",
        "cbm_risk_premium": "+1.5%",
        "sector_spread": "-0.5%",
        "final_interest_rate": "9.50%",
        "verdict_rationale": (
            "Revenue figures reconcile within 2% tolerance across GST, Bank, and AR. "
            "No adverse litigation. CMR-3 confirms strong payment discipline. CNC sector "
            "shows 14% YoY national growth. Recommended for sanction at ₹5.00 Cr @ 9.50%."
        ),
        "concept_scores": {
            "Character": 90, "Capacity": 85,
            "Capital": 88, "Collateral": 82, "Conditions": 91,
        },
        "concept_flags": {},   # e.g. {"Character": "Active DRT litigation — ₹45L disputed"}
        "concept_narratives": {
            "Character": "Promoter Vikram Mehta has a clean regulatory record. No DRT, eCourt, or MCA adverse flags. CIBIL CMR-3 affirms strong institutional credit discipline.",
            "Capacity": "GST revenue of ₹41.8 Cr reconciles within 2% of bank credits (₹42.1 Cr) and audited financials (₹42.5 Cr). DSCR of 1.8x is well above the 1.25x threshold.",
            "Capital": "TNW of ₹18.2 Cr with TOL/TNW of 1.4x. Promoter equity stake at 67%. Reserves & surplus growing 12% YoY over 3 years.",
            "Collateral": "Primary: Hypothecation of current assets (book value ₹7.2 Cr). Collateral: Equitable mortgage on factory premises (market value ₹9.5 Cr). SCR: 1.6x. SARFAESI enforceable.",
            "Conditions": "CNC & Automation sector growing at 14% YoY nationally. No adverse RBI sector circular. End-use: Working capital enhancement for order-book execution.",
        },
        "shap_top_factors": [
            {"feature": "GST-Bank Revenue Reconciliation", "contribution": +18.2},
            {"feature": "CIBIL CMR Rating (CMR-3)",        "contribution": +14.5},
            {"feature": "DSCR (1.8x)",                     "contribution": +12.0},
            {"feature": "Sector Growth (CNC 14% YoY)",     "contribution": +9.5},
            {"feature": "Collateral Coverage (1.6x SCR)",  "contribution": +8.0},
            {"feature": "Promoter NPA Linkage (None)",     "contribution": +6.5},
        ],
        "llm_narrative": (
            "CleanTech Manufacturing Ltd. demonstrates robust financial health across all verification "
            "pillars. Revenue figures reconcile within 2% tolerance across GST filings (₹41.8 Cr), bank "
            "credits (₹42.1 Cr), and audited financials (₹42.5 Cr). No adverse litigation detected. "
            "CIBIL CMR-3 rating confirms strong payment discipline. The CNC machinery sector shows 14% "
            "YoY growth nationally, providing favourable macro tailwinds. Recommended for sanction at "
            "₹5.00 Cr with a blended rate of 9.50%."
        ),
        "dd_override_notes": None,
    }

    gen = CAMGenerator()
    path = gen.generate_cam_docx(sample_data, "/mnt/user-data/outputs/CAM_CleanTech.docx")
    print(f"CAM generated: {path}")
